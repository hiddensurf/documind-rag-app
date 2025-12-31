"""
Document Loader - Load and parse various document types including CAD and 3D files
"""
import logging
from pathlib import Path
from typing import List
import fitz  # PyMuPDF
import pymupdf4llm
from docx import Document as DocxDocument
from llama_index.core import Document
from app.utils.cad_loader import CADLoader
from app.utils.stl_loader import STLLoader
from unstructured.partition.auto import partition

logger = logging.getLogger(__name__)

class DocumentLoader:
    """Load documents from various file formats"""
    
    def __init__(self):
        self.cad_loader = CADLoader()
        self.stl_loader = STLLoader()
    
    def load_document(self, file_path: str, file_id: str = None, file_name: str = None) -> List[Document]:
        """
        Load document from file
        
        Args:
            file_path: Path to the file
            file_id: Optional file ID for metadata
            file_name: Optional original filename
            
        Returns:
            List of Document objects
        """
        file_ext = Path(file_path).suffix.lower()
        
        # CAD files (DWG/DXF)
        if file_ext in ['.dwg', '.dxf']:
            return self._load_cad(file_path, file_id, file_name)
        
        # 3D model files (STL)
        elif file_ext == '.stl':
            return self._load_stl(file_path, file_id, file_name)
        
        # PDF files
        elif file_ext == '.pdf':
            return self._load_pdf(file_path, file_id, file_name)
        
        # DOCX files
        elif file_ext == '.docx':
            return self._load_docx(file_path, file_id, file_name)
        
        # Plain text files
        elif file_ext in ['.txt', '.md']:
            return self._load_text(file_path, file_id, file_name)
        
        # Try Unstructured for other formats
        else:
            return self._load_with_unstructured(file_path, file_id, file_name)
    
    def _load_cad(self, file_path: str, file_id: str, file_name: str) -> List[Document]:
        """Load CAD file (DWG/DXF)"""
        logger.info(f"Loading CAD: {file_path}")
        
        result = self.cad_loader.load_cad_file(file_path, file_id, file_name)
        
        if not result["success"]:
            logger.warning(f"CAD loading failed: {result['error']}")
            return [Document(
                text=result["text_content"],
                metadata={
                    "doc_id": file_id,
                    "file_name": file_name,
                    "file_type": "cad",
                    "conversion_status": "failed",
                    "error": result["error"]
                }
            )]
        
        return result["documents"]
    
    def _load_stl(self, file_path: str, file_id: str, file_name: str) -> List[Document]:
        """Load STL 3D model file"""
        logger.info(f"Loading STL: {file_path}")
        
        result = self.stl_loader.load_stl(file_path, file_id, file_name)
        
        if not result["success"]:
            logger.warning(f"STL loading failed: {result['error']}")
            return [Document(
                text=result["text_content"],
                metadata={
                    "doc_id": file_id,
                    "file_name": file_name,
                    "file_type": "stl",
                    "error": result["error"]
                }
            )]
        
        return result["documents"]
    
    def _load_pdf(self, file_path: str, file_id: str = None, file_name: str = None) -> List[Document]:
        """Load PDF file using PyMuPDF"""
        logger.info(f"Loading PDF: {file_path}")
        
        try:
            md_text = pymupdf4llm.to_markdown(file_path)
            
            doc = fitz.open(file_path)
            num_pages = len(doc)
            doc.close()
            
            logger.info(f"Extracted {len(md_text)} characters from {num_pages} pages")
            
            metadata = {
                "file_type": "pdf",
                "num_pages": num_pages
            }
            
            if file_id:
                metadata["doc_id"] = file_id
            if file_name:
                metadata["file_name"] = file_name
            
            return [Document(text=md_text, metadata=metadata)]
            
        except Exception as e:
            logger.error(f"Error loading PDF: {str(e)}")
            raise
    
    def _load_docx(self, file_path: str, file_id: str = None, file_name: str = None) -> List[Document]:
        """Load DOCX file"""
        logger.info(f"Loading DOCX: {file_path}")
        
        try:
            doc = DocxDocument(file_path)
            
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            text = "\n\n".join(paragraphs)
            
            logger.info(f"Extracted {len(text)} characters from {len(paragraphs)} paragraphs")
            
            metadata = {
                "file_type": "docx",
                "num_paragraphs": len(paragraphs)
            }
            
            if file_id:
                metadata["doc_id"] = file_id
            if file_name:
                metadata["file_name"] = file_name
            
            return [Document(text=text, metadata=metadata)]
            
        except Exception as e:
            logger.error(f"Error loading DOCX: {str(e)}")
            raise
    
    def _load_text(self, file_path: str, file_id: str = None, file_name: str = None) -> List[Document]:
        """Load plain text file"""
        logger.info(f"Loading text file: {file_path}")
        
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                text = f.read()
            
            logger.info(f"Loaded {len(text)} characters")
            
            file_ext = Path(file_path).suffix.lower()
            
            metadata = {
                "file_type": file_ext.lstrip('.')
            }
            
            if file_id:
                metadata["doc_id"] = file_id
            if file_name:
                metadata["file_name"] = file_name
            
            return [Document(text=text, metadata=metadata)]
            
        except Exception as e:
            logger.error(f"Error loading text file: {str(e)}")
            raise
    
    def _load_with_unstructured(self, file_path: str, file_id: str = None, file_name: str = None) -> List[Document]:
        """Load file using Unstructured library (supports many formats)"""
        logger.info(f"Loading with Unstructured: {file_path}")
        
        try:
            from unstructured.partition.auto import partition
            
            # Partition the document
            elements = partition(filename=file_path)
            
            # Extract text from elements
            text_parts = [str(el) for el in elements]
            text = "\n\n".join(text_parts)
            
            logger.info(f"Extracted {len(text)} characters using Unstructured")
            
            file_ext = Path(file_path).suffix.lower()
            
            metadata = {
                "file_type": file_ext.lstrip('.'),
                "loader": "unstructured",
                "element_count": len(elements)
            }
            
            if file_id:
                metadata["doc_id"] = file_id
            if file_name:
                metadata["file_name"] = file_name
            
            return [Document(text=text, metadata=metadata)]
            
        except Exception as e:
            logger.error(f"Error loading with Unstructured: {str(e)}")
            raise ValueError(f"Unsupported file type or error loading: {str(e)}")

# Create singleton instance
document_loader = DocumentLoader()